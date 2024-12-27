import asyncio
import io
import json
import logging
import random
from abc import abstractmethod
from datetime import datetime
from enum import Enum
from time import sleep
from typing import Dict, Optional, Any, Callable, Iterable, AsyncIterable
import gitlab
import base64

import langchain_core.documents
from atlassian import Jira, Confluence
from gitlab.v4.objects import Project, ProjectFile
from docx.api import Document as DocxDocument

from ai_core.base import ComponentType, create_tool_name
from ai_core.data_source.model.document import Document, DocumentEncoder
from ai_core.data_source.utils.opensearch_utils import switch_to_new_index, IndexingStatistics, asearch, \
    delete_indices_with_alias
from ai_core.data_source.utils.utils import create_data_source_id, bs4_extractor, truncate_content, get_first
from ai_core.data_source.utils.time_utils import get_iso_8601_current_time, DATE_FORMAT
from ai_core.data_source.utils.opensearch_utils import create_opensearch_index_name
from ai_core.data_source.collection import Collection
from ai_core.data_source.utils.jira_utils import filter_required_fields, cleanse_fields
from ai_core.data_source.vectorstore.search_type import Similarity

from langchain_core.retrievers import BaseRetriever
from langchain.tools.retriever import create_retriever_tool
from langchain.tools import Tool
from langchain_community.document_loaders import RecursiveUrlLoader, UnstructuredPDFLoader, UnstructuredFileLoader

from opensearchpy import OpenSearch, AsyncOpenSearch

from pydantic import BaseModel, Field, ConfigDict


logger = logging.getLogger(__name__)


class DataSourceType(Enum):
    TEXT = "text"
    PDF_FILE = "pdf_file"
    CONFLUENCE = "confluence"
    GITLAB = "gitlab"
    GITLAB_DISCUSSION = "gitlab_discussion"
    URL = "url"
    DOC_FILE = "doc_file"
    JIRA = "jira"


class DataSource(BaseModel):
    model_config = ConfigDict(arbitrary_types_allowed=True)

    id: str = Field(str, frozen=True)
    name: str = Field(str, frozen=True)
    description: str
    data_source_type: DataSourceType
    collections: Dict[str, Collection] = Field(default_factory=dict)
    opensearch_hosts: str
    opensearch_auth: tuple[str, str]
    opensearch_client: OpenSearch = None
    async_opensearch_client: AsyncOpenSearch = None
    preview_data: str = None
    save_data_in_progress: bool = Field(default=False)

    PREVIEW_DATA_MAX_LENGTH: int = 10000

    def model_post_init(self, __context: Any) -> None:
        opensearch_client_args = {
            "hosts": self.opensearch_hosts,
            "http_compress": True,
            "http_auth": self.opensearch_auth,
            "use_ssl": True,
            "verify_certs": False,
            "ssl_show_warn": False,
            "auto_create_index": False
        }

        self.opensearch_client = OpenSearch(**opensearch_client_args)
        self.async_opensearch_client = AsyncOpenSearch(**opensearch_client_args)

    def add_collection(self, collection_name: str, llm_api_provider: str, llm_api_key: str, llm_api_url: str,
                       llm_embedding_model_name: str, last_update_succeeded_at: Optional[datetime] = None):

        collection = Collection(
            datasource_id=self.id,
            name=collection_name,
            llm_api_provider=llm_api_provider,
            llm_api_key=llm_api_key,
            llm_api_url=llm_api_url,
            llm_embedding_model_name=llm_embedding_model_name,
            vectorstore_hosts=self.opensearch_hosts,
            vectorstore_auth=self.opensearch_auth,
            last_update_succeeded_at=last_update_succeeded_at,
        )

        self.collections[collection.name] = collection

        return collection

    def get_latest_collection(self) -> Collection:
        sorted_collections = (
            sorted(
                filter(lambda c: c.last_update_succeeded_at is not None, self.collections.values()),
                key=lambda collection: collection.last_update_succeeded_at))

        if sorted_collections:
            return sorted_collections[-1]
        else:
            logger.warning(f"No collection found with last_updated_succeeded_at for data source {self.id}")

            return None

    def as_retriever(self, search_type=Similarity(k=4)) -> BaseRetriever:
        latest_collection = self.get_latest_collection()

        '''
            Retriever의 메타데이터 입니다. 
            특정 instance를 식별하기 위한 용도 등으로 사용합니다.
        '''
        metadata = {
            "data_source_id": latest_collection.datasource_id,
            "collection_name": latest_collection.name,
            "llm_api_provider": latest_collection.llm_api_provider,
            "llm_api_url": latest_collection.llm_api_url,
            "llm_embedding_model": latest_collection.llm_embedding_model.name,
            "created_at": get_iso_8601_current_time()
        }

        index_alias = latest_collection.name
        search_kwargs = search_type.search_kwargs
        search_kwargs["index_name"] = index_alias

        return latest_collection.vectorstore.as_retriever(
            search_type=search_type.name, metadata=metadata, search_kwargs=search_kwargs)

    def as_retriever_tool(self, name: str, description: str, search_type=Similarity(k=4)) -> Tool:
        if not name:
            raise ValueError("name must be provided")
        if not description:
            raise ValueError("description must be provided")

        return create_retriever_tool(
            self.as_retriever(search_type=search_type),
            name=name,
            description=description
        )

    @abstractmethod
    def _download_data(self, **kwargs) -> Iterable[Document]:
        pass

    @abstractmethod
    def load_preview_data(self, **kwargs) -> str:
        pass

    async def cancel_save_data_task(self, save_task: asyncio.Task) -> None:
        '''
        데이터 저장 작업을 취소합니다.
        '''

        save_task.cancel()
        try:
            await save_task
        except asyncio.CancelledError:
            logger.info(f"Save task of datasource {self.id} was cancelled.")

    def save_data(self, last_update_succeeded_at: str, **kwargs) -> dict:
        logger.info(f"Start saving data of datasource {self.id}")

        documents = self._download_data(**kwargs)

        result = self._save_data(documents=documents, last_update_succeeded_at=last_update_succeeded_at, kwargs=kwargs)

        logger.info(f"Finish saving data of datasource {self.id}")

        return result

    def _save_data(self, documents: Iterable[Document], last_update_succeeded_at: str, **kwargs) -> dict:
        '''
        docx, pdf file, confluence 등 다양한 소스로부터 불러온 텍스트 데이터를 opensearch에 저장합니다.

        :param kwargs: DataSource의 구현체마다 각각 다른 파라메터를 전달 받습니다.
        :return result dict: index_name, index_alias, doc_count, doc_total_bytes, doc_max_bytes, doc_min_bytes, doc_avg_bytes
        '''

        if self.save_data_in_progress:
            raise ValueError("Data saving is already in progress.")

        try:
            self.save_data_in_progress = True

            index_alias = self.id
            index_alias_in_progress = f"{index_alias}_in_progress"
            index_name = create_opensearch_index_name(self.id)

            # 이전에 실패했던 작업이 있을 경우, 생성된 index를 삭제합니다.
            delete_indices_with_alias(self.opensearch_client, index_alias_in_progress)

            prev_index_names = []
            if self.opensearch_client.indices.exists_alias(name=index_alias):
                prev_index_names = self.opensearch_client.indices.get_alias(name=index_alias).keys()
            result = {"index_name": index_name, "index_alias": index_alias}

            # Opensearch index 생성
            logger.info(f"Create opensearch index: {index_name}")
            self.opensearch_client.indices.create(index=index_name)
            self.opensearch_client.indices.put_alias(index=index_name, name=index_alias_in_progress)

            # 문서 다운로드
            logger.info(f"Download data from {self.data_source_type.name}. {kwargs}")

            # 문서 색인
            logger.info(f"Indexing documents to opensearch index: {index_name}")
            indexing_statistics = IndexingStatistics()
            for document in documents:
                doc_id = None
                if document.metadata:
                    doc_id = document.metadata.get("doc_id", None)
                    document.metadata["last_update_succeeded_at"] = last_update_succeeded_at
                    document.metadata["data_source_type"] = self.data_source_type.value

                json_string = json.dumps(document, cls=DocumentEncoder, indent=2, ensure_ascii=False)

                self.opensearch_client.index(index=index_name, body=json_string, id=doc_id)

                indexing_statistics.increment(len(json_string.encode('utf-8')))

            stats = indexing_statistics.values()
            logger.info(f"Indexed {stats['doc_count']} documents to opensearch index: {index_name}. "
                        f"stats: {stats}")

            logger.info(f"Refresh opensearch index: {index_name}")
            self.opensearch_client.indices.refresh(index=index_name)

            # alias 업데이트
            logger.info(f"Update opensearch index alias: {index_alias}, index name: {index_name}")
            switch_to_new_index(self.opensearch_client, index_alias, index_alias_in_progress, index_name,
                                prev_index_names)

            return {**result, **stats}
        except Exception as e:
            logger.error(f"Error occurred while saving data: {e}")
            raise e
        finally:
            self.save_data_in_progress = False

    def _update_data(self, **kwargs) -> dict:
        return self.save_data(**kwargs)

    def update_data(self, **kwargs) -> dict:
        '''
        데이터를 업데이트합니다.
        :param kwargs: DataSource의 구현체마다 각각 다른 파라메터를 전달 받습니다.
        :return result dict: index_name, index_alias, doc_count, doc_total_bytes, doc_max_bytes, doc_min_bytes, doc_avg_bytes
        '''

        index_alias = self.id
        if self.opensearch_client.indices.exists_alias(name=index_alias):
            return self._update_data(**kwargs)
        else:
            raise ValueError(f"Index alias {index_alias} does not exist.")

    async def read_data(self, since: str = None) -> AsyncIterable[Document]:
        '''
        opensearch에서 데이터를 읽어옵니다.
        '''

        logger.info(f"Read data from opensearch index: {self.id}")
        index_alias = self.id
        client = self.async_opensearch_client
        query = {"match_all": {}}

        if since:
            query = {
                "range": {
                    "metadata.last_update_succeeded_at": {
                        "gte": since
                    }
                }
            }

        return asearch(async_opensearch_client=client, index_name=index_alias, query=query)

    async def read_preview_data(self) -> str:
        index_alias = self.id
        client = self.async_opensearch_client

        if self.preview_data:
            return self.preview_data
        if not await client.indices.exists(index=index_alias):
            self.preview_data = ""
            return self.preview_data

        query = {
            "size": 1,
            "from": 0,
            "query": {
                "match_all": {}
            }
        }

        result = await self.async_opensearch_client.search(
            index=index_alias,
            body=query,
            size=1
        )

        if result["hits"]["hits"]:
            content = result["hits"]["hits"][0]["_source"]["content"]
            self.preview_data = truncate_content(content, self.PREVIEW_DATA_MAX_LENGTH)
        else:
            self.preview_data = ""

        return self.preview_data

    def delete_data(self) -> None:
        client = self.opensearch_client
        index_alias = self.id
        index_names = []

        if client.indices.exists_alias(name=index_alias):
            index_names = client.indices.get_alias(name=index_alias).keys()

        for index_name in index_names:
            logger.info(f"Delete opensearch index: {index_name}")
            client.indices.delete(index=index_name)

    async def adelete_data(self) -> None:
        client = self.async_opensearch_client
        index_alias = self.id
        index_names = []

        if await client.indices.exists_alias(name=index_alias):
            index_names = (await client.indices.get_alias(name=index_alias)).keys()

        for index_name in index_names:
            logger.info(f"Delete opensearch index: {index_name}")
            await client.indices.delete(index=index_name)


class TextDataSource(DataSource):
    def _download_data(self, raw_text: Iterable[str]) -> Iterable[Document]:
        for content in raw_text:
            yield Document(content=content, metadata={})

    def load_preview_data(self, raw_text: Iterable[str]) -> str:
        return truncate_content(get_first(raw_text), self.PREVIEW_DATA_MAX_LENGTH)


class PdfFileDataSource(DataSource):
    def _create_metadata(self, doc_file_path: str):
        file_name = doc_file_path.split("/")[-1]

        return {"file_name": file_name} if file_name else {}

    def _download_data(self, pdf_file_path: str) -> Iterable[Document]:
        loader = UnstructuredPDFLoader(pdf_file_path)
        docs = loader.lazy_load()

        for doc in docs:
            yield Document(content=doc.page_content, metadata=self._create_metadata(pdf_file_path))

    def load_preview_data(self, pdf_file_path: str) -> str:
        from PyPDF2 import PdfReader
        def read_first_page_pdf(file_path, encoding='utf-8'):
            with open(file_path, 'rb') as file:
                reader = PdfReader(io.BytesIO(file.read()))
                first_page = reader.pages[0]
                return first_page.extract_text().encode('utf-8').decode(encoding)

        first_page_content = read_first_page_pdf(pdf_file_path)

        if first_page_content:
            return truncate_content(first_page_content, self.PREVIEW_DATA_MAX_LENGTH)
        else:
            return ""


class ConfluenceDataSource(DataSource):
    def _create_metadata(self, confluence, page):
        metadata = {}

        if page.get("id"):
            metadata["id"] = page["id"]
        if page.get("title"):
            metadata["title"] = page["title"]
        if page.get("_links") and page["_links"].get("webui"):
            webui_link = page["_links"]["webui"]
            webui_link = webui_link[1:] if webui_link.startswith("/") else webui_link
            metadata["url"] = f"{confluence.url}{webui_link}"

        return metadata if metadata else {}

    # page는 한 번에 최대 100개씩만 가져와진다.
    def get_pages(self, confluence: Confluence, space_key: str, start: int, batch_size: int):
        return confluence.get_all_pages_from_space(
            space=space_key, start=start, limit=batch_size, status=None, expand="body.view")

    def get_all_pages(self,
                      confluence: Confluence,
                      space_key: str,
                      batch_size: int = 100,
                      request_interval: int = 0.25) -> Iterable[Document]:
        logger.info(f"Start scrape confluence pages from space_key: {space_key}.")
        start = 0
        total_page_count = 0

        while True:
            pages = self.get_pages(confluence, space_key, start, batch_size)
            documents = []

            for page in pages:
                content = page["body"]["view"]["value"]
                metadata = self._create_metadata(confluence, page)
                documents.append(Document(content=content, metadata=metadata))

            page_count = len(documents)
            total_page_count += page_count
            logger.info(f"{page_count} confluence pages are scraped from {start} to {start + page_count - 1}")

            yield from documents

            if page_count < batch_size:
                break

            start += batch_size
            sleep(request_interval)

        logger.info(f"Finish scrape confluence pages from space: {space_key}. Total {total_page_count} pages are scraped.")

    '''
    TDE Confluence는 초당 5회 이내로 API 호출을 제한합니다.

    :param url: Confluence 서버 URL
    :param user_name: Confluence 사용자 이름
    :param access_token: Confluence API 토큰
    :param batch_size: 한 번에 가져올 페이지 수. 최대 100개까지 가능합니다.
    '''
    def _download_data(self, url: str, access_token: str, space_key: str, batch_size: int = 100) -> Iterable[Document]:
        confluence = Confluence(url=url, token=access_token)
        return self.get_all_pages(confluence, space_key, batch_size)

    def load_preview_data(self, url: str, access_token: str, space_key: str, batch_size: int = 1) -> str:
        confluence = Confluence(url=url, token=access_token)
        pages = self.get_pages(confluence, space_key, start=0, batch_size=batch_size)
        first_page = get_first(pages)

        if first_page:
            content = first_page.get("body", "").get("view", "").get("value", "")
            return truncate_content(content, self.PREVIEW_DATA_MAX_LENGTH)
        else:
            return ""


class GitlabDataSource(DataSource):
    def _create_metadata(self, file_content: ProjectFile):
        metadata = {}
        if file_content.file_path:
            metadata["file_path"] = file_content.file_path
            metadata["doc_id"] = base64.standard_b64encode(file_content.file_path.encode('utf-8')).decode('utf-8')
        if file_content.last_commit_id:
            metadata["last_commit_id"] = file_content.last_commit_id

        return metadata

    def _download_data(self, url: str, namespace: str, project_name: str, branch: str, private_token: str) \
            -> Iterable[Document]:
        def filter_valid_files(tree):
            include_extensions = [".txt", ".md", ".rst", ".html", ".htm", ".xml", ".json", ".csv", ".tsv", ".yaml",
                                  ".yml", ".log", ".py", ".java", ".js", ".ts", ".c", ".cpp", ".h", ".hpp", ".cs",
                                  ".go", ".php", ".rb", ".sh", ".scala", ".sql", ".ipynb", ".conf", ".properties"]

            for item in tree:
                if item['type'] == 'blob' and any([item['path'].endswith(ext) for ext in include_extensions]):
                    yield item

        gl = gitlab.Gitlab(url=url, private_token=private_token)
        project = gl.projects.get(f"{namespace}/{project_name}")
        tree = project.repository_tree(ref=branch, recursive=True, all=True)

        for item in filter_valid_files(tree):
            file_content = project.files.get(file_path=item['path'], ref=branch)
            try:
                decoded_content = base64.b64decode(file_content.content).decode('utf-8')
                metadata = self._create_metadata(file_content)
                yield Document(content=decoded_content, metadata=metadata)
            except UnicodeDecodeError:
                logger.error(f"Can't decode {item['path']}")

    def _get_random_file_from_commits(self, project, branch, max_commits=5):
        try:
            commits = project.commits.list(ref_name=branch, all=True)
            file_paths = set()

            for commit in commits[:max_commits]:
                diff = project.commits.get(commit.id).diff()
                for change in diff:
                    file_paths.add(change['new_path'])

                if len(file_paths) >= 50:  # 충분한 파일을 수집했다면 중단
                    break

            if not file_paths:
                return None

            return random.choice(list(file_paths))

        except gitlab.exceptions.GitlabError as e:
            logger.error(f"Error getting file list: {e}")
            return None

    def load_preview_data(self, url: str, namespace: str, project_name: str, branch: str, private_token: str) -> str:
        gl = gitlab.Gitlab(url=url, private_token=private_token)
        project = gl.projects.get(f"{namespace}/{project_name}")
        preview_file_path = self._get_random_file_from_commits(project, branch)
        file_content = project.files.get(file_path=preview_file_path, ref=branch)
        try:
            decoded_content = base64.b64decode(file_content.content).decode('utf-8')
            return truncate_content(decoded_content, self.PREVIEW_DATA_MAX_LENGTH)
        except UnicodeDecodeError:
            logger.error(f"Can't decode {preview_file_path}")

    def _get_file_paths_for_update(self, gitlab_project_key: str, gitlab_project: Project, branch: str,
                                   since: str) -> (set, set):
        '''
        :param gitlab_project_key: gitlab project key
        :param gitlab_project: gitlab project object
        :param branch: gitlab branch name
        :param since: ex) ""2024-06-01T09:00:00.000+0900""
        :return: update_file_paths, delete_file_paths
        '''

        commits = gitlab_project.commits.list(ref_name=branch, since=since, get_all=True)
        commits = sorted(commits, key=lambda x: x.committed_date)

        upsert_file_paths = set()
        delete_file_paths = set()

        logger.info(f"Getting file list of gitlab project {gitlab_project_key}, branch {branch} to update")

        for commit in commits:
            diff = gitlab_project.commits.get(commit.id).diff(get_all=True)
            for changed_file in diff:
                if changed_file.get("deleted_file"):
                    delete_file_paths.add(changed_file.get("old_path"))
                    upsert_file_paths.discard(changed_file.get("old_path"))
                elif changed_file.get("renamed_file"):
                    upsert_file_paths.discard(changed_file.get("old_path"))
                    upsert_file_paths.add(changed_file.get("new_path"))
                    delete_file_paths.add(changed_file.get("old_path"))
                else:
                    # new_file, modified_file
                    upsert_file_paths.add(changed_file.get("new_path"))
                    delete_file_paths.discard(changed_file.get("new_path"))

        logger.info(f"There are {len(upsert_file_paths)} added or changed files of gitlab project {gitlab_project_key}, "
              f"branch {branch} since {since}")

        logger.info(f"There are {len(delete_file_paths)} deleted files of gitlab project {gitlab_project_key}, branch {branch} "
              f"since {since}")

        return upsert_file_paths, delete_file_paths

    def _download_data_for_update(self, gitlab_project: Project, branch: str, upsert_file_paths: set) \
            -> Iterable[Document]:
        '''
        :param gitlab_project: gitlab project object
        :param branch: gitlab branch name
        :param upsert_file_paths: set of file paths to update
        :return: result: iterable documents
        '''

        for upsert_file_path in upsert_file_paths:
            try:
                file_for_upsert: ProjectFile = gitlab_project.files.get(file_path=upsert_file_path, ref=branch)
                decoded_content = base64.b64decode(file_for_upsert.content).decode('utf-8')
                yield Document(content=decoded_content, metadata=self._create_metadata(file_for_upsert))
            except gitlab.exceptions.GitlabGetError:
                # subtree는 파일을 가져오지 못함
                logger.error(f"Error getting file from gitlab: {upsert_file_path}")

    def _update_documents(self, project_key: str, gitlab_project: Project, branch: str,
                          documents: Iterable[Document], delete_file_paths: set, last_update_succeeded_at: str) -> dict:
        '''
        :param url: gitlab url
        :param namespace: gitlab namespace
        :param project_name: gitlab project name
        :param branch: gitlab branch name
        :param private_token: gitlab private token
        :param last_update_succeeded_at: ex) ""2024-06-01T09:00:00.000+0900""
        :return: result dict: index_alias, doc_count, doc_total_bytes, doc_max_bytes, doc_min_bytes, doc_avg_bytes
        '''

        indexing_statistics = IndexingStatistics()

        logger.info(f"Start updating data of gitlab project {project_key}, branch {branch}")
        index_alias = self.id
        for document in documents:
            document.metadata["last_update_succeeded_at"] = last_update_succeeded_at
            json_string = json.dumps(document, cls=DocumentEncoder, indent=2, ensure_ascii=False)
            doc_id = document.metadata.get("doc_id")
            self.opensearch_client.index(index=index_alias, body=json_string, id=doc_id)

            indexing_statistics.increment(len(json_string.encode('utf-8')))

        for delete_file_path in delete_file_paths:
            file_for_delete: ProjectFile = gitlab_project.files.get(file_path=delete_file_path, ref=branch)
            self.opensearch_client.delete(index=index_alias, id=file_for_delete.blob_id)

            indexing_statistics.decrement(len(file_for_delete.content.encode('utf-8')))

        logger.info(f"Finish updating data of gitlab project {project_key}, branch {branch}")
        result = {"index_alias": index_alias}

        return {**result, **indexing_statistics.values()}

    def _update_data(self, url: str, namespace: str, project_name: str, branch: str, 
                     private_token: str, since: str, last_update_succeeded_at: str) -> dict:
        '''
        :param url: gitlab url
        :param namespace: gitlab namespace
        :param project_name: gitlab project name
        :param branch: gitlab branch name
        :param private_token: gitlab private token
        :param since: 이 시각 이후에 커밋된 파일을 업데이트합니다. ex) ""2024-06-02T09:00:00.000+0900""
        :param last_update_succeeded_at: 마지막 업데이트 시각을 저장합니다. ex) ""2024-06-01T09:00:00.000+0900""
        :return: result dict: index_alias, doc_count, doc_total_bytes, doc_max_bytes, doc_min_bytes, doc_avg_bytes
        '''

        project_key = f"{namespace}/{project_name}"
        gl = gitlab.Gitlab(url=url, private_token=private_token)
        project = gl.projects.get(project_key)
        
        upsert_file_paths, delete_file_paths = (
            self._get_file_paths_for_update(project_key, project, branch, since))
        documents = self._download_data_for_update(project, branch, upsert_file_paths)
        result = self._update_documents(project_key, project, branch, documents, delete_file_paths, last_update_succeeded_at)

        return result


class UrlDataSource(DataSource):
    def _create_metadata(self, document: langchain_core.documents.Document):
        url = document.metadata.get("source")
        return {"url": url} if url else {}

    def _download_data(self, url: str, max_depth: int, base_url: str, extractor: Callable[[str], str] = bs4_extractor) \
            -> Iterable[Document]:
        '''
        :param url: scraping 하고자 하는 웹사이트의 root url
        :param max_depth: 재귀적으로 링크를 타고 들어갈 최대 깊이
        :param base_url: prevent_outside=True 일 경우 base_url을 기준으로 외부 링크를 제외합니다.
        :param extractor: 페이지에서 텍스트를 추출하는 함수
        :return: generator[str]
        '''

        loader = RecursiveUrlLoader(url=url, max_depth=max_depth, extractor=extractor, base_url=base_url)
        docs = loader.lazy_load()

        for doc in docs:
            yield Document(content=doc.page_content, metadata=self._create_metadata(doc))

    def load_preview_data(self, url: str, base_url: str, extractor: Callable[[str], str] = bs4_extractor) -> str:
        loader = RecursiveUrlLoader(url=url, max_depth=1, extractor=extractor, base_url=base_url)
        docs = loader.lazy_load()

        first_doc = get_first(docs)
        if first_doc:
            return truncate_content(first_doc.page_content, self.PREVIEW_DATA_MAX_LENGTH)
        else:
            return ""


class DocumentFileDataSource(DataSource):
    def _create_metadata(self, doc_file_path: str):
        file_name = doc_file_path.split("/")[-1]

        return {"file_name": file_name} if file_name else {}

    def _download_data(self, doc_file_path: str) -> list[str]:
        loader = UnstructuredFileLoader(file_path=doc_file_path, mode="single")
        docs = loader.lazy_load()

        for doc in docs:
            yield Document(content=doc.page_content, metadata=self._create_metadata(doc_file_path))

    def load_preview_data(self, doc_file_path: str) -> str:
        doc = DocxDocument(doc_file_path)
        first_page_content = []

        for paragraph in doc.paragraphs:
            if paragraph.paragraph_format.page_break_before:
                break
            first_page_content.append(paragraph.text)

        first_page = '\n'.join(first_page_content)

        if first_page:
            return truncate_content(first_page, self.PREVIEW_DATA_MAX_LENGTH)
        else:
            return ""


class JiraDataSource(DataSource):
    def _create_metadata(self, cleaned_fields: dict) -> dict:
        metadata_keys = ["key", "updated", "assignee"]
        metadata = {key: cleaned_fields[key] for key in metadata_keys if key in cleaned_fields}
        key = cleaned_fields.get("key")
        if key:
            metadata["doc_id"] = key
            metadata.pop("key")

        return metadata if metadata else {}

    def _download_data(self, url: str, project_key: str, access_token: str, start: int = 0, limit: int = 1000,
                       call_interval: float = 0.25) -> Iterable[Document]:
        jira = Jira(url=url, token=access_token)
        issue_keys = jira.get_project_issuekey_all(project_key, start=start, limit=limit)

        while issue_keys:
            logger.info(f"Retrieved {len(issue_keys)} issue keys of {project_key} project. start from {issue_keys[0]}")

            for issue_key in issue_keys:
                issue = jira.issue(issue_key)

                required_fields = filter_required_fields(issue)
                cleaned_fields = cleanse_fields(required_fields)

                content = json.dumps(cleaned_fields, indent=2, ensure_ascii=False)
                metadata = self._create_metadata(cleaned_fields)

                yield Document(content=content, metadata=metadata)

                start += 1
                sleep(call_interval)

            issue_keys = jira.get_project_issuekey_all(project_key, start=start, limit=limit)
        else:
            logger.info(f"No issue keys retrieved. Finish loading jira issues for {project_key} project.")

    def load_preview_data(self, url: str, project_key: str, access_token: str, start: int = 0, limit: int = 1) -> str:
        jira = Jira(url=url, token=access_token)
        issue_keys = jira.get_project_issuekey_all(project_key, start=start, limit=limit)

        if issue_keys:
            issue = jira.issue(issue_keys[0])
            required_fields = filter_required_fields(issue)
            cleaned_fields = cleanse_fields(required_fields)

            return json.dumps(cleaned_fields, indent=2, ensure_ascii=False)
        else:
            return ""

    def _download_data_for_update(self, url: str, project_key: str, access_token: str, since: str,
                                  call_interval: float) -> Iterable[Document]:
        '''
        :param url: jira url
        :param project_key: jira project key
        :param access_token: jira access token
        :param since: ex) "2024-06-01T00:00:00.000+0000"
        :return: iterable documents
        '''

        start_date = datetime.strptime(since, DATE_FORMAT).strftime("%Y-%m-%d %H:%M")
        jira = Jira(url=url, token=access_token)
        issues = jira.jql(f"project = '{project_key}' AND updated >= '{start_date}'")
        logger.info(f"Retrieved {len(issues.get('issues', []))} issues updated since {since}")

        for issue in issues.get("issues", []):
            required_fields = filter_required_fields(issue)
            cleaned_fields = cleanse_fields(required_fields)

            content = json.dumps(cleaned_fields, indent=2, ensure_ascii=False)
            metadata = self._create_metadata(cleaned_fields)

            yield Document(content=content, metadata=metadata)

            sleep(call_interval)

    def _update_data(self, url: str, project_key: str, access_token: str, since: str, last_update_succeeded_at: str,
    call_interval: float = 0.25) -> dict:
        '''
        신규 업데이트 된 이슈를 가져와서 색인합니다.
        last_update_succeeded_at 이후에 업데이트 된 이슈만 가져옵니다.

        :param url: jira url
        :param project_key: jira project key
        :param access_token: jira access token
        :param since: 이 시간 이후에 업데이트된 jira issue를 업데이트 합니다. ex) "2024-06-01T00:00:00.000+0000"
        :param last_update_succeeded_at: 마지막 업데이트한 시각을 저장합니다. ex) "2024-06-02T00:00:00.000+0000"
        :return: result dict: index_alias, doc_count, doc_total_bytes, doc_max_bytes, doc_min_bytes, doc_avg_bytes
        '''

        logger.info(f"Start updating data of jira project {project_key}")
        index_alias = self.id

        documents = self._download_data_for_update(
            url, project_key, access_token, since, call_interval)
        indexing_statistics = IndexingStatistics()

        logger.info(f"Indexing documents to opensearch index: {index_alias}")
        for document in documents:
            document.metadata["last_update_succeeded_at"] = last_update_succeeded_at
            body = json.dumps(document, cls=DocumentEncoder, indent=2, ensure_ascii=False)
            doc_id = document.metadata.get("doc_id")
            self.opensearch_client.index(index=index_alias, body=body, id=doc_id)

            indexing_statistics.increment(len(body.encode('utf-8')))

        indexing_result = indexing_statistics.values()

        logger.info(f"Indexed {indexing_result['doc_count']} documents to opensearch index: {self.id}. ")
        logger.info(f"Finish updating data of jira project {project_key}")

        return {"index_alias": index_alias, **indexing_result}


class GitlabDiscussionDataSource(DataSource):
    def _get_file_content(self, project: Project, file_path: str, sha: str):
        """특정 파일의 내용 가져오기"""
        # 파일 내용 가져오기 (base64로 인코딩되어 있음)
        try:
            file_obj = project.files.get(file_path=file_path, ref=sha)

            content = file_obj.decode()

            return content.decode('utf-8')
        except Exception as e:
            return None

    def _get_code_context(self, file_content: str, start_line: int, end_line: int):
        """코드 컨텍스트 출력"""
        lines = file_content.split('\n')
        context = lines[max(0, start_line - 1):min(len(lines), end_line)]

        texts = "\n"
        for i, line in enumerate(context, start=start_line):
            texts += f"{i}: {line}\n"

        return texts

    def _download_data(self, url: str, namespace: str, project_name: str, private_token: str):
        gl = gitlab.Gitlab(url=url, private_token=private_token)
        project = gl.projects.get(f"{namespace}/{project_name}")
        merge_requests = list(project.mergerequests.list(all=True))

        discussion_texts = []
        for mr in merge_requests:
            # merge_header = f"\nMerge Request #{mr.iid}: {mr.title}"
            discussions = list(mr.discussions.list(all=True))

            for discussion in discussions:
                note = discussion.attributes['notes'][0]
                if note['system']:
                    continue

                author = note['author']['name']
                if author == "DevSecOps IT보안운영(임시)":
                    continue

                header = f"Discussion ID: {discussion.id}"
                texts = ""
                if 'position' in note:
                    position = note['position']
                    file_path = position['new_path']
                    start_line = position.get('new_line')
                    if start_line is None:
                        continue

                    start_line -= 5 if start_line > 5 else start_line
                    end_line = start_line + 10  # 컨텍스트로 5줄 표시

                    texts += f"\n    파일: {file_path}"
                    file_content = self._get_file_content(project, file_path, mr.sha)
                    if file_content:
                        texts += f"\n    코드 컨텍스트:"
                        texts += self._get_code_context(file_content, start_line, end_line)

                for note in discussion.attributes['notes']:
                    texts += f"\n    작성자: {note['author']['name']}"
                    texts += f"\n    내용: {note['body']}"

                if texts:
                    discussion_texts.append(header + texts)

        return discussion_texts

    def load_preview_data(self, url: str, namespace: str, project_name: str, private_token: str):
        discussion_texts: Iterable[Document] = self._download_data(url, namespace, project_name, private_token)

        first_element = next(iter(discussion_texts))
        if first_element:
            return truncate_content(first_element.content, self.PREVIEW_DATA_MAX_LENGTH)


def create_data_source(data_source_name: str, created_by: str, description: str, data_source_type: str,
                       opensearch_hosts: str, opensearch_auth: tuple[str, str]) \
        -> DataSource:
    datasource_data = {
        "id": create_data_source_id(created_by, data_source_name),
        "name": data_source_name,
        "description": description,
        "data_source_type": data_source_type,
        "opensearch_hosts": opensearch_hosts,
        "opensearch_auth": opensearch_auth
    }

    if data_source_type == DataSourceType.TEXT.value:
        return TextDataSource(**datasource_data)
    elif data_source_type == DataSourceType.PDF_FILE.value:
        return PdfFileDataSource(**datasource_data)
    elif data_source_type == DataSourceType.CONFLUENCE.value:
        return ConfluenceDataSource(**datasource_data)
    elif data_source_type == DataSourceType.GITLAB.value:
        return GitlabDataSource(**datasource_data)
    elif data_source_type == DataSourceType.URL.value:
        return UrlDataSource(**datasource_data)
    elif data_source_type == DataSourceType.DOC_FILE.value:
        return DocumentFileDataSource(**datasource_data)
    elif data_source_type == DataSourceType.JIRA.value:
        return JiraDataSource(**datasource_data)
    elif data_source_type == DataSourceType.GITLAB_DISCUSSION.value:
        return GitlabDiscussionDataSource(**datasource_data)
    else:
        raise ValueError(f"Invalid data source type: {data_source_type}")


def create_data_source_tool(name: str, username: str, datasource: DataSource) -> Tool:
    tool_name = create_tool_name(ComponentType.DATASOURCE, name, username)
    return datasource.as_retriever_tool(tool_name, datasource.description)